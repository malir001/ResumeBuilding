from docx import Document
from docx.shared import Pt

def build_resume(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.add_heading(data['name'], level=1)
    doc.add_paragraph(f"{data['email']} | {data['phone']} | {data['linkedin']} | {data['github']}")

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(data['summary'])

    doc.add_heading("Technical Skills", level=2)
    doc.add_paragraph(", ".join(data['skills']))

    doc.add_heading("Work Experience", level=2)
    for exp in data['experience']:
        doc.add_paragraph(f"{exp['role']} – {exp['company']} ({exp['duration']})", style='List Bullet')
        for responsibility in exp['responsibilities']:
            doc.add_paragraph(f"• {responsibility}", style='List Bullet 2')

    doc.add_heading("Education", level=2)
    doc.add_paragraph(f"{data['education_degree']} – {data['education_institute']} ({data['education_year']})")

    return doc
